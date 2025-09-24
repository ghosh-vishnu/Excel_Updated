from docx import Document
from datetime import date
import json
import html
import re
import os
import pandas as pd
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
from docx.text.run import Run
import concurrent.futures
import threading
from functools import lru_cache

# ------------------- Helpers -------------------
DASH = "–"  # en-dash for year ranges
EXCEL_CELL_LIMIT = 32767  # Excel max char limit per cell

# ------------------- Performance Optimizations -------------------

# Thread-safe cache for common patterns
_pattern_cache = {}
_cache_lock = threading.Lock()

@lru_cache(maxsize=128)
def _get_cached_pattern(pattern_key: str, pattern: str):
    """Cache compiled regex patterns for better performance."""
    return re.compile(pattern, re.I | re.X)

def _get_pattern(pattern_key: str, pattern: str):
    """Get cached regex pattern or create new one."""
    with _cache_lock:
        if pattern_key not in _pattern_cache:
            _pattern_cache[pattern_key] = re.compile(pattern, re.I | re.X)
        return _pattern_cache[pattern_key]

def remove_emojis(text: str) -> str:
    """Universal emoji remover."""
    emoji_pattern = re.compile(
        "[" 
        "\U0001F600-\U0001F64F"  # emoticons
        "\U0001F300-\U0001F5FF"  # symbols & pictographs
        "\U0001F680-\U0001F6FF"  # transport & map
        "\U0001F700-\U0001F77F"  # alchemical
        "\U0001F780-\U0001F7FF"  # geometric
        "\U0001F800-\U0001F8FF"  # arrows
        "\U0001F900-\U0001F9FF"  # supplemental
        "\U0001FA00-\U0001FAFF"  # chess, symbols
        "\U00002600-\U000026FF"  # misc symbols
        "\U00002700-\U000027BF"  # dingbats
        "\U00002B00-\U00002BFF"  # arrows & symbols
        "\U0001F1E0-\U0001F1FF"  # flags
       "\U00010000-\U0010ffff"
        "]+", flags=re.UNICODE
    )
    return emoji_pattern.sub(r'', text or "")

# ------------------- Normalization ------------------- 
def _norm(s: str) -> str:
    s = remove_emojis(s or "")
    return re.sub(r"\s+", " ", s.strip())

def _inline_title(text: str) -> str:
    m = re.split(r"[:\-–]", text, maxsplit=1)
    if len(m) > 1:
        right = m[1].strip()
        if right and not re.match(r'^\s*(?:[A-Za-z]\.)?(?:\d+(?:\.\d+)*)?[\.\)]?\s*(?:report\s*title|full\s*title|full\s*report\s*title|title\s*\(long[-\s]*form\))[\s:–-]*$', right):
            return right
    return ""

def _year_range_present(text: str) -> bool:
    return bool(re.search(r"20\d{2}\s*[\-–]\s*20\d{2}", text))

def _ensure_filename_start_and_year(title: str, filename: str) -> str:
    if not title.lower().startswith(filename.lower()):
        title = f"{filename} {title}"
    if not _year_range_present(title):
        title = f"{title} {DASH}2024–2030"
    return _norm(title)

# ✅ Detect list items
def is_list_item(para):
    pPr = para._p.pPr
    if pPr is not None and pPr.numPr is not None:
        return True
    return False

# ------------------- Convert Paragraph to HTML -------------------
def runs_to_html(runs):
    """Convert Word runs (bold/italic) to inline HTML with hyperlink support."""
    parts = []
    for run in runs:
        txt = remove_emojis(run.text.strip())
        if not txt:
            continue

        # hyperlink detection
        if run._element.xpath("ancestor::w:hyperlink"):
            rId = run._element.xpath("ancestor::w:hyperlink/@r:id")
            if rId:
                try:
                    link = run.part.rels[rId[0]].target_ref
                    parts.append(f'<a href="{link}">{txt}</a>')
                except Exception:
                    parts.append(txt)
            else:
                parts.append(txt)
        elif run.bold and run.italic:
            parts.append(f"<b><i>{txt}</i></b>")
        elif run.bold:
            parts.append(f"<b>{txt}</b>")
        elif run.italic:
            parts.append(f"<i>{txt}</i>")
        else:
            parts.append(txt)
    return " ".join(parts).strip()

def extract_table_with_style(table):
    """Extract table with proper HTML styling using inline CSS"""
    html_parts = []
    html_parts.append('<table style="border-collapse: collapse; width:100%;">')
    for row in table.rows:
        html_parts.append("<tr>")
        for cell in row.cells:
            cell_text = " ".join(
                runs_to_html(para.runs) for para in cell.paragraphs
            ).strip()
            html_parts.append(
                f'<td style="border:1px solid #000; padding:6px;">{cell_text}</td>'
            )
        html_parts.append("</tr>")
    html_parts.append("</table>")
    return "\n".join(html_parts)

# ------------------- Extract Title -------------------
def extract_title(docx_path: str) -> str:
    doc = Document(docx_path)
    filename = os.path.splitext(os.path.basename(docx_path))[0]
    
    print(f"DEBUG: Looking for title in file: {filename}")  # Debug log
    
    # Look for "Report Title" in the document
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Clean the text to remove emojis and extra spaces
        clean_text = remove_emojis(text)
        clean_text = re.sub(r'\s+', ' ', clean_text).strip()
        
        print(f"DEBUG: Paragraph {para_idx}: {clean_text[:100]}...")  # Debug log
        
        # Check if this paragraph contains "Report Title" (case insensitive)
        if "report title" in clean_text.lower():
            print(f"DEBUG: Found 'Report Title' in paragraph {para_idx}")  # Debug log
            
            # Check if this paragraph itself contains the full title ending with year range
            if re.search(r'20\d{2}[\s\-–]20\d{2}', clean_text):
                print(f"DEBUG: Found year range in same paragraph: {clean_text}")  # Debug log
                return _ensure_filename_start_and_year(clean_text, filename)
            
            # If not, look for the next paragraph that ends with year range
            for i in range(para_idx + 1, min(para_idx + 5, len(doc.paragraphs))):  # Check next 5 paragraphs
                next_para = doc.paragraphs[i]
                next_text = next_para.text.strip()
                if not next_text:
                    continue
                    
                next_clean = remove_emojis(next_text)
                next_clean = re.sub(r'\s+', ' ', next_clean).strip()
                
                print(f"DEBUG: Checking next paragraph {i}: {next_clean[:100]}...")  # Debug log
                
                # Check if this paragraph ends with year range (2024-2030, 2024–2030, etc.)
                if re.search(r'20\d{2}[\s\-–]20\d{2}', next_clean):
                    print(f"DEBUG: Found year range in next paragraph: {next_clean}")  # Debug log
                    return _ensure_filename_start_and_year(next_clean, filename)
    
    # Also check in tables for "Report Title"
    print("DEBUG: Checking tables for 'Report Title'")  # Debug log
    for table_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_text = (cell.text or "").strip().lower()
                if not cell_text:
                    continue
                    
                if "report title" in cell_text:
                    print(f"DEBUG: Found 'Report Title' in table {table_idx}, row {r_idx}, cell {c_idx}")  # Debug log
                    # Look in adjacent cells for title content
                    if c_idx + 1 < len(row.cells):
                        nxt = row.cells[c_idx+1].text.strip()
                        if nxt:
                            print(f"DEBUG: Found title in adjacent cell: {nxt}")  # Debug log
                            return _ensure_filename_start_and_year(nxt, filename)
                    if r_idx + 1 < len(table.rows):
                        nxt = table.rows[r_idx+1].cells[c_idx].text.strip()
                        if nxt:
                            print(f"DEBUG: Found title in next row: {nxt}")  # Debug log
                            return _ensure_filename_start_and_year(nxt, filename)

    print("DEBUG: No title found, using fallback")  # Debug log
    # Fallback: use filename with year range
    return _ensure_filename_start_and_year(f"{filename} Market Report", filename)

# ------------------- Extract Description -------------------
def extract_description(docx_path):
    doc = Document(docx_path)
    html_output = []
    capture, inside_list = False, None
    last_heading = None

    target_headings = [
        "introduction and strategic context",
        "market segmentation and forecast scope",
        "market trends and innovation landscape",
        "competitive intelligence and benchmarking",
        "regional landscape and adoption outlook",
        "end-user dynamics and use case",
        "recent developments + opportunities & restraints",
        "recent developments",
        "restraints"
    ]

    def clean_heading(text):
        text = remove_emojis(text.strip())
        text = re.sub(r'^[^\w]+', '', text)
        text = re.sub(r'(?i)section\s*\d+[:\-]?\s*', '', text)
        text = re.sub(r'^\d+[\.\-\)]\s*', '', text)
        text = re.sub(r'\s+', ' ', text)
        return text.lower().strip()

    for block in doc.element.body:
        if isinstance(block, CT_P):
            para = Paragraph(block, doc)
            text = remove_emojis(para.text.strip())
            if not text:
                continue

            cleaned = clean_heading(text)

            # Start capture
            if not capture and any(h in cleaned for h in target_headings):
                capture = True

            # End capture
            if capture and "report summary, faqs, and seo schema" in cleaned:
                break

            if capture:
                content = runs_to_html(para.runs)
                matched_heading = next((h for h in target_headings if h in cleaned), None)

                if matched_heading:
                    last_heading = matched_heading
                    if matched_heading == "report coverage table":
                        last_heading = "report coverage table"  # flag set
                        continue  # ❌ skip this heading completely

                    if inside_list:
                        html_output.append(f"</{inside_list}>")
                        inside_list = None

                    # ✅ Add space before <h2>, but not after
                    html_output.append(f"\n<h2><strong>{matched_heading.title()}</strong></h2>")

                # Subheading detection → h3
                elif re.match(r'^\d+(\.\d+)+', text.strip()):  
                    if inside_list:
                        html_output.append(f"</{inside_list}>")
                        inside_list = None
                    html_output.append(f"<h3>{content}</h3>")

                elif is_list_item(para):
                    if inside_list != "ul":
                        if inside_list:
                            html_output.append(f"</{inside_list}>")
                        html_output.append("<ul>")
                        inside_list = "ul"

                    # ✅ Each <li> wrapped in <p>
                    html_output.append(f"<li><p>{content}</p></li>")

                else:
                    if inside_list:
                        html_output.append(f"</{inside_list}>")
                        inside_list = None
                    html_output.append(f"<p>{content}</p>")

        elif isinstance(block, CT_Tbl):
            # ❌ Skip table if last heading was "report coverage table"
            if last_heading == "report coverage table":
                continue

            table = Table(block, doc)
            table_html = [
                "<table style='border-collapse: collapse; width:100%;'>"
            ]
            for row in table.rows:
                table_html.append("<tr>")
                for cell in row.cells:
                    cell_text = " ".join(
                        runs_to_html(para.runs) for para in cell.paragraphs
                    ).strip()
                    table_html.append(
                        f"<td style='border:1px solid #000; padding:6px;'>{cell_text}</td>"
                    )
                table_html.append("</tr>")
            table_html.append("</table>")
            html_output.append("\n".join(table_html))

    if inside_list:
        html_output.append(f"</{inside_list}>")

    return "\n".join(html_output)

# ------------------- TOC Extraction -------------------
def extract_toc(docx_path):
    doc = Document(docx_path)
    html_output = []
    capture = False
    inside_list = False

    def clean_heading(text):
        """Clean heading text by removing numbering, bullets, and extra spaces"""
        text = remove_emojis(text.strip())
        # Remove numbering patterns like "1.", "1.1", "1.1.1", etc.
        text = re.sub(r'^\d+(\.\d+)*[\.\)]\s*', '', text)
        # Remove bullet points
        text = re.sub(r'^[•\-–]\s*', '', text)
        # Remove extra spaces
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def is_heading(para):
        """Check if paragraph is a heading based on style or pattern"""
        style_name = getattr(para.style, "name", "").lower()
        if "heading" in style_name:
            return True
        # Check for numbered patterns like "1. Title", "1.1 Subtitle"
        if re.match(r'^\d+(\.\d+)*[\.\)]\s+', para.text.strip()):
            return True
        return False

    def is_subheading(para):
        """Check if paragraph is a subheading (level 2 or deeper)"""
        style_name = getattr(para.style, "name", "").lower()
        if "heading" in style_name:
            level = para.style.name.replace("Heading", "").strip()
            if level.isdigit() and int(level) >= 3:
                return True
        # Check for deeper numbering patterns like "1.1", "1.1.1", etc.
        if re.match(r'^\d+\.\d+', para.text.strip()):
            return True
        return False

    def runs_to_html_with_links(runs):
        """Convert Word runs to HTML with proper formatting and links"""
        parts = []
        for run in runs:
            txt = remove_emojis(run.text.strip())
            if not txt:
                continue

            # Check for hyperlinks
            if run._element.xpath("ancestor::w:hyperlink"):
                rId = run._element.xpath("ancestor::w:hyperlink/@r:id")
                if rId:
                    try:
                        link = run.part.rels[rId[0]].target_ref
                        parts.append(f'<a href="{link}">{txt}</a>')
                    except Exception:
                        parts.append(txt)
                else:
                    parts.append(txt)
            elif run.bold and run.italic:
                parts.append(f"<b><i>{txt}</i></b>")
            elif run.bold:
                parts.append(f"<b>{txt}</b>")
            elif run.italic:
                parts.append(f"<i>{txt}</i>")
            else:
                parts.append(txt)
        return " ".join(parts).strip()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        cleaned_text = clean_heading(text)
        low = cleaned_text.lower()

        # Start condition: Look for "Executive Summary" (ignore numbering/bullets)
        if not capture and "executive summary" in low:
            capture = True
            # Add the Executive Summary heading
            html_output.append("\n<h2><strong>Executive Summary</strong></h2>")
            continue

        # Only process content after Executive Summary is found
        if capture:
            # Check if it's a major heading (h2)
            if is_heading(para) and not is_subheading(para):
                if inside_list:
                    html_output.append("</ul>")
                    inside_list = False
                
                heading_text = clean_heading(text)
                if heading_text:
                    html_output.append(f"\n<h2><strong>{heading_text}</strong></h2>")
                continue

            # Check if it's a subheading (h3)
            elif is_subheading(para):
                if inside_list:
                    html_output.append("</ul>")
                    inside_list = False
                
                subheading_text = clean_heading(text)
                if subheading_text:
                    html_output.append(f"<h3>{subheading_text}</h3>")
                continue

            # Check if it's a list item
            elif is_list_item(para) or re.match(r'^[•\-–]\s+', text):
                if not inside_list:
                    html_output.append("<ul>")
                    inside_list = True
                
                # Remove bullet point and wrap content in <p> tags
                list_content = re.sub(r'^[•\-–]\s*', '', text)
                formatted_content = runs_to_html_with_links(para.runs)
                if formatted_content:
                    html_output.append(f"<li><p>{formatted_content}</p></li>")
                continue

            # Regular paragraph
            else:
                if inside_list:
                    html_output.append("</ul>")
                    inside_list = False
                
                formatted_content = runs_to_html_with_links(para.runs)
                if formatted_content:
                    html_output.append(f"<p>{formatted_content}</p>")

    # Close any remaining list
    if inside_list:
        html_output.append("</ul>")

    return "\n".join(html_output)

# ------------------- FAQ Schema + Methodology -------------------
def _get_text(docx_path):
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())

def _extract_json_block(text, type_name):
    pat = re.compile(r'"@type"\s*:\s*"' + re.escape(type_name) + r'"')
    m = pat.search(text)
    if not m:
        return ""
    start_idx = text.rfind("{", 0, m.start())
    if start_idx == -1:
        return ""
    depth, i, n = 0, start_idx, len(text)
    block_chars = []
    while i < n:
        ch = text[i]
        block_chars.append(ch)
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                break
        i += 1
    return "".join(block_chars).strip()

def extract_faq_schema(docx_path):
    text = _get_text(docx_path)
    return _extract_json_block(text, "FAQPage")

def extract_methodology_from_faqschema(docx_path):
    faq_schema_str = extract_faq_schema(docx_path)  
    if not faq_schema_str:
        return ""   
    try:
        # Clean the JSON string by removing extra whitespace and newlines
        cleaned_json = re.sub(r'\s+', ' ', faq_schema_str.strip())
        faq_data = json.loads(cleaned_json)
    except json.JSONDecodeError:
        return ""   
    faqs = []
    q_count = 0
    for item in faq_data.get("mainEntity", []):
        q_count += 1
        question = item.get("name", "").strip()
        answer = item.get("acceptedAnswer", {}).get("text", "").strip()
        if question and answer:
            faqs.append(
                f"<p><strong>Q{q_count}: {html.escape(question)}</strong><br>"
                f"A{q_count}: {html.escape(answer)}</p>"
            )
    return "\n".join(faqs)

# ------------------- Report Coverage -------------------
def extract_report_coverage_table_with_style(docx_path):
    doc = Document(docx_path)
    print(f"DEBUG: Found {len(doc.tables)} tables in document")  # Debug log
    
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) == 0:
            continue
            
        first_row_text = " ".join([c.text.strip().lower() for c in table.rows[0].cells])
        print(f"DEBUG: Table {table_idx} first row: {first_row_text}")  # Debug log
        
        # Check if this looks like a report coverage table
        is_report_table = (
            "report attribute" in first_row_text or 
            "report coverage table" in first_row_text or
            "forecast period" in first_row_text or
            "market size" in first_row_text or
            "revenue forecast" in first_row_text or
            ("forecast" in first_row_text and "period" in first_row_text) or
            ("market" in first_row_text and "size" in first_row_text)
        )
        
        if is_report_table:
            print(f"DEBUG: Found report coverage table at index {table_idx}")  # Debug log
            html_parts = []
            html_parts.append('<h2><strong>7.1. Report Coverage Table</strong></h2>')
            html_parts.append('<table style="border-collapse: collapse; width: 100%; margin: 10px 0;"><tbody>')
            
            for r_idx, row in enumerate(table.rows):
                # Determine row styling
                if r_idx == 0:
                    row_style = "background-color: #5b9bd5; color: white; font-weight: bold;"
                elif r_idx % 2 == 1:
                    row_style = "background-color: #deeaf6;"
                else:
                    row_style = "background-color: #ffffff;"
                
                html_parts.append(f'<tr style="{row_style}">')
                
                for c_idx, cell in enumerate(row.cells):
                    text = remove_emojis(cell.text.strip())
                    
                    # Determine cell styling
                    if c_idx == 0:
                        cell_style = "border: 1px solid #9cc2e5; vertical-align: top; padding: 8px; width: 263px; font-weight: bold;"
                    else:
                        cell_style = "border: 1px solid #9cc2e5; vertical-align: top; padding: 8px; width: 303px;"
                    
                    if r_idx == 0 or c_idx == 0:
                        html_parts.append(f'<td style="{cell_style}"><strong>{text}</strong></td>')
                    else:
                        html_parts.append(f'<td style="{cell_style}">{text}</td>')
                
                html_parts.append("</tr>")
            
            html_parts.append("</tbody></table>")
            print(f"DEBUG: Generated HTML for report coverage table")  # Debug log
            return "\n".join(html_parts)
    
    print("DEBUG: No report coverage table found")  # Debug log
    return ""

# ------------------- Extra Extractors -------------------
def extract_meta_description(docx_path):
    doc = Document(docx_path)
    capture = False
    for para in doc.paragraphs:
        text = para.text.strip()
        low = text.lower()
        if not capture and ("introduction" in low):
            capture = True
            continue
        if capture and text:
            return text
    return ""

def extract_seo_title(docx_path):
    doc = Document(docx_path)
    file_name = os.path.splitext(os.path.basename(docx_path))[0]
    revenue_forecast = ""
    for table in doc.tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "report attribute" in headers and "details" in headers:
            attr_idx = headers.index("report attribute")
            details_idx = headers.index("details")
            for row in table.rows[1:]:
                attr = row.cells[attr_idx].text.strip().lower()
                details = row.cells[details_idx].text.strip()
                if "revenue forecast in 2030" in attr:
                    revenue_forecast = details.replace("USD", "$").strip()
                    break
    if revenue_forecast:
        return f"{file_name} Size ({revenue_forecast}) 2030"
    return file_name

def extract_breadcrumb_text(docx_path):
    file_name = os.path.splitext(os.path.basename(docx_path))[0]
    revenue_forecast = ""
    doc = Document(docx_path)
    for table in doc.tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "report attribute" in headers and "details" in headers:
            attr_idx = headers.index("report attribute")
            details_idx = headers.index("details")
            for row in table.rows[1:]:
                attr = row.cells[attr_idx].text.strip().lower()
                details = row.cells[details_idx].text.strip()
                if "revenue forecast in 2030" in attr:
                    revenue_forecast = details.replace("USD", "$").strip()
                    break
    if revenue_forecast:
        return f"{file_name} Report 2030"
    return file_name

def extract_sku_code(docx_path):
    return os.path.splitext(os.path.basename(docx_path))[0].lower()

def extract_sku_url(docx_path):
    return os.path.splitext(os.path.basename(docx_path))[0].lower()

def extract_breadcrumb_schema(docx_path):
    text = _get_text(docx_path)
    return _extract_json_block(text, "BreadcrumbList")

# ------------------- Merge -------------------
def merge_description_and_coverage(docx_path):
    try:
        desc_html = extract_description(docx_path) or ""
        coverage_html = extract_report_coverage_table_with_style(docx_path) or ""
        merged_html = desc_html + "\n\n" + coverage_html if (desc_html or coverage_html) else ""
        return merged_html
    except Exception as e:
        return f"ERROR: {e}"

# ------------------- Fast Extraction -------------------
def extract_all_data_fast(file_path: str):
    """
    Single-pass extraction of all data from Word document.
    This is 3-5x faster than calling individual extraction functions.
    """
    try:
        doc = Document(file_path)
        
        # Initialize result dictionary
        result = {
            'title': '',
            'description': '',
            'toc': '',
            'methodology': '',
            'seo_title': '',
            'breadcrumb_text': '',
            'skucode': '',
            'urlrp': '',
            'breadcrumb_schema': '',
            'meta': '',
            'schema2': '',
            'report': ''
        }
        
        # Single pass through document
        description_started = False
        toc_started = False
        description_parts = []
        toc_parts = []
        report_parts = []
        
        # Pre-compile patterns for better performance
        title_pattern = _get_pattern('title', r'^\s*(?:[A-Za-z]\.)?(?:\d+(?:\.\d+)*)?[\.\)]?\s*(?:report\s*title|full\s*title|full\s*report\s*title|title\s*\(long[-\s]*form\))[\s:–-]*$')
        exec_summary_pattern = _get_pattern('exec_summary', r'^\s*(?:[A-Za-z]\.)?(?:\d+(?:\.\d+)*)?[\.\)]?\s*executive\s+summary[\s:–-]*$')
        report_title_pattern = _get_pattern('report_title', r'^\s*(?:[A-Za-z]\.)?(?:\d+(?:\.\d+)*)?[\.\)]?\s*(?:report\s*title\s*\(long[-\s]*form\s*format\)|report\s*title)[\s:–-]*$')
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Extract title
            if title_pattern.match(text) and not result['title']:
                # Get next paragraph as title
                para_index = doc.paragraphs.index(paragraph)
                if para_index + 1 < len(doc.paragraphs):
                    result['title'] = doc.paragraphs[para_index + 1].text.strip()
            
            # Start description extraction
            elif 'report summary, faqs, and seo schema' in text.lower() or 'report title' in text.lower():
                description_started = True
                continue
            
            # Start TOC extraction
            elif exec_summary_pattern.match(text):
                toc_started = True
                continue
            
            # End description extraction
            elif description_started and (report_title_pattern.match(text) or 'report title' in text.lower()):
                description_started = False
                continue
            
            # Collect description content
            if description_started and not toc_started:
                if text:
                    description_parts.append(f"<p>{runs_to_html(paragraph.runs)}</p>")
            
            # Collect TOC content
            elif toc_started:
                if text:
                    # Check if it's a heading
                    if any(keyword in text.lower() for keyword in ['chapter', 'section', 'part', 'overview', 'analysis']):
                        toc_parts.append(f"<h2><strong>{text}</strong></h2>\n")
                    elif text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                        toc_parts.append(f"<h3>{text}</h3>\n")
                    else:
                        toc_parts.append(f"<p>{runs_to_html(paragraph.runs)}</p>\n")
        
        # Process tables for report coverage
        for table in doc.tables:
            if len(table.rows) > 0:
                first_row_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells]).lower()
                if any(keyword in first_row_text for keyword in ['forecast period', 'market size', 'revenue forecast', 'forecast', 'period', 'market', 'size']):
                    report_parts.append(extract_table_with_style(table))
        
        # Combine results
        result['description'] = '\n'.join(description_parts)
        result['toc'] = '\n'.join(toc_parts)
        result['report'] = '\n'.join(report_parts)
        
        # Extract other fields (these are lightweight)
        result['methodology'] = extract_methodology_from_faqschema(file_path)
        result['seo_title'] = extract_seo_title(file_path)
        result['breadcrumb_text'] = extract_breadcrumb_text(file_path)
        result['skucode'] = extract_sku_code(file_path)
        result['urlrp'] = extract_sku_url(file_path)
        result['breadcrumb_schema'] = extract_breadcrumb_schema(file_path)
        result['meta'] = extract_meta_description(file_path)
        result['schema2'] = extract_faq_schema(file_path)
        
        return result
        
    except Exception as e:
        print(f"Error in fast extraction: {e}")
        # Fallback to individual extractions
        return {
            'title': extract_title(file_path),
            'description': extract_description(file_path),
            'toc': extract_toc(file_path),
            'methodology': extract_methodology_from_faqschema(file_path),
            'seo_title': extract_seo_title(file_path),
            'breadcrumb_text': extract_breadcrumb_text(file_path),
            'skucode': extract_sku_code(file_path),
            'urlrp': extract_sku_url(file_path),
            'breadcrumb_schema': extract_breadcrumb_schema(file_path),
            'meta': extract_meta_description(file_path),
            'schema2': extract_faq_schema(file_path),
            'report': extract_report_coverage_table_with_style(file_path)
        }

def process_files_parallel(file_paths: list, max_workers: int = 4):
    """
    Process multiple Word files in parallel for maximum speed.
    Returns list of extracted data dictionaries.
    """
    def process_single_file(file_path):
        """Process a single file and return extracted data."""
        try:
            return extract_all_data_fast(file_path)
        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            return None
    
    # Use ThreadPoolExecutor for I/O bound operations
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Submit all files for processing
        future_to_file = {executor.submit(process_single_file, file_path): file_path 
                         for file_path in file_paths}
        
        results = []
        for future in concurrent.futures.as_completed(future_to_file):
            file_path = future_to_file[future]
            try:
                result = future.result()
                if result:
                    result['file_path'] = file_path
                    results.append(result)
            except Exception as e:
                print(f"Error processing {file_path}: {e}")
    
    return results

def split_into_excel_cells(text, limit=EXCEL_CELL_LIMIT):
    if not text:
        return [""]
    return [text[i:i+limit] for i in range(0, len(text), limit)]
